#!/usr/bin/env python3
"""
perla-textify — converts "compatible but not plain-text" document formats
into something Perla's model can actually look at.

Two different strategies, chosen per format:

  - TEXT extraction (DOCX, XLSX, CSV/TSV, Jupyter notebooks, ODT): these
    are fundamentally text/data documents where layout doesn't carry much
    meaning, so their content is pulled out as plain text and inlined
    into the prompt the same way perla-companion.py already inlines
    .py/.md/.nix uploads.

  - PAGE-IMAGE rasterization (PDF, PPTX): these are fundamentally VISUAL
    documents — diagrams, charts, slide layout, embedded images, table
    formatting — all of which a text-only extraction throws away. These
    are instead rendered page-by-page/slide-by-slide into PNGs and
    handed back as a list of image bytes, which perla-companion.py sends
    through the SAME path as any other user-attached image (mimo-v2.5's
    vision input), not through the text-inlining path at all.

Why a separate module rather than folding this into perla-companion.py:
  - The daemon's core job (session management, HTTP, MCP proxying) has
    nothing to do with document parsing, and each format needs its own
    (optional) third-party library/binary. Keeping conversion in one
    file means a missing/broken dependency for one format never risks
    breaking the daemon's import at startup, and this file can be
    unit-tested or swapped independently.
  - Every heavy import below is LAZY (done inside the function that needs
    it, not at module load) specifically so that companion still starts
    and every OTHER upload path keeps working even if, say, poppler
    isn't installed on this machine — only .pdf/.pptx uploads would
    fail, with a clear "install X" message, not a daemon-wide crash.

Public entrypoints:
    extract_text(raw_bytes, filename)   -> (text, error)
    rasterize_pages(raw_bytes, filename) -> (list_of_png_bytes, error)
    convert(raw_bytes, filename)        -> ("text"|"images", payload, error)

`convert` is what perla-companion.py should actually call — it looks up
which strategy a given extension uses and returns a tagged result, so the
caller doesn't need its own per-format routing table duplicating the one
here.

This module does NOT touch the network, or any Perla-specific state —
it's a pure bytes-in/converted-out module, testable on its own with
`python3 perla-textify.py <file>` (text formats print to stdout; PDF/PPTX
write numbered PNGs next to the input file instead).
"""

import csv
import io
import json
import os
import re
import shutil
import subprocess
import tempfile
import zipfile


# ---------------------------------------------------------------------------
# Per-format size ceilings. Independent from perla-companion.py's
# MAX_TEXT_UPLOAD_BYTES (which gates the raw upload size before it ever
# reaches this module) — these instead cap the size of the EXTRACTED
# text, since a 500KB PDF can expand into much more raw text than a
# 500KB source file would, and a huge extraction is exactly as capable
# of blowing the model's context window as a huge plain-text upload.
MAX_EXTRACTED_CHARS = 200_000  # ~50-60k tokens, a generous but bounded slice
TRUNCATION_NOTE = (
    "\n\n[... extracted text truncated at {limit} characters — "
    "the original file is longer than this excerpt ...]"
)


def _truncate(text):
    if len(text) <= MAX_EXTRACTED_CHARS:
        return text
    return text[:MAX_EXTRACTED_CHARS] + TRUNCATION_NOTE.format(limit=MAX_EXTRACTED_CHARS)


# ---------------------------------------------------------------------------
# Format extractors. Each takes raw bytes and returns (text, error).
# Naming convention: _extract_<ext without dot>.
# ---------------------------------------------------------------------------

# ---------------------------------------------------------------------------
# Page-image rasterization (PDF, PPTX) — these formats are handed back as
# a list of PNG byte-blobs instead of extracted text. Both share the same
# per-page/per-image limits so a huge deck or PDF can't blow the model's
# per-message image budget.
# ---------------------------------------------------------------------------
MAX_PAGES = 20  # matches a generous but bounded slide-deck/report length
RASTER_DPI = 150  # good text legibility without producing huge PNGs


def _png_bytes_from_pil(img):
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    return buf.getvalue()


def _rasterize_pdf(raw):
    """Render a PDF's pages to PNG bytes, one per page. Returns
    (list_of_png_bytes, error). Truncates at MAX_PAGES rather than
    failing outright on a long document — the caller can tell the user
    only the first N pages were attached."""
    try:
        from pdf2image import convert_from_bytes
    except ImportError:
        return None, (
            "PDF page rendering isn't installed on this machine "
            "(pip install pdf2image --break-system-packages, plus the "
            "poppler-utils system package)"
        )

    try:
        images = convert_from_bytes(
            raw, dpi=RASTER_DPI, fmt="png",
            first_page=1, last_page=MAX_PAGES,
        )
    except Exception as e:
        # pdf2image raises PDFPageCountError/PDFSyntaxError/etc depending
        # on the poppler failure; the message text is what actually
        # matters to surface, the exact exception class doesn't.
        msg = str(e)
        if "Unable to get page count" in msg or "poppler" in msg.lower():
            return None, (
                "couldn't render this PDF — poppler-utils may not be "
                f"installed on this machine ({msg})"
            )
        if "password" in msg.lower() or "encrypted" in msg.lower():
            return None, "PDF is password-protected — can't render it"
        return None, f"couldn't render this PDF: {msg}"

    if not images:
        return None, "PDF has no pages to render"

    pages = [_png_bytes_from_pil(img) for img in images]
    return pages, None


def _rasterize_pptx(raw):
    """Render a PPTX's slides to PNG bytes, one per slide. There is no
    pure-Python slide renderer worth depending on (python-pptx only reads
    the XML, it doesn't draw anything) — instead this shells out to
    LibreOffice headless to convert PPTX -> PDF, then reuses
    _rasterize_pdf on the result, so the actual rendering logic exists in
    exactly one place. Requires `soffice`/`libreoffice` on PATH."""
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        return None, (
            "PPTX rendering needs LibreOffice installed on this machine "
            "(the 'soffice' or 'libreoffice' command isn't on PATH)"
        )

    with tempfile.TemporaryDirectory(prefix="perla-pptx-") as tmpdir:
        src_path = os.path.join(tmpdir, "input.pptx")
        with open(src_path, "wb") as f:
            f.write(raw)

        try:
            result = subprocess.run(
                [soffice, "--headless", "--norestore", "--convert-to", "pdf",
                 "--outdir", tmpdir, src_path],
                capture_output=True, timeout=90,
            )
        except subprocess.TimeoutExpired:
            return None, "converting this PPTX to a renderable form timed out"
        except Exception as e:
            return None, f"couldn't run LibreOffice to convert this PPTX: {e}"

        pdf_path = os.path.join(tmpdir, "input.pdf")
        if result.returncode != 0 or not os.path.exists(pdf_path):
            stderr = result.stderr.decode(errors="replace").strip()
            return None, f"LibreOffice couldn't convert this PPTX to a renderable form: {stderr[:200]}"

        with open(pdf_path, "rb") as f:
            pdf_bytes = f.read()

    return _rasterize_pdf(pdf_bytes)


    try:
        import docx  # python-docx
    except ImportError:
        return None, (
            "DOCX support isn't installed on this machine "
            "(pip install python-docx --break-system-packages)"
        )
    try:
        document = docx.Document(io.BytesIO(raw))
    except Exception as e:
        return None, f"couldn't open as a DOCX: {e}"

def _extract_docx(raw):
    try:
        import docx  # python-docx
    except ImportError:
        return None, (
            "DOCX support isn't installed on this machine "
            "(pip install python-docx --break-system-packages)"
        )
    try:
        document = docx.Document(io.BytesIO(raw))
    except Exception as e:
        return None, f"couldn't open as a DOCX: {e}"

    parts = []
    for para in document.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))

    text = "\n".join(parts)
    if not text.strip():
        return None, "no extractable text found in this DOCX (it may be empty)"
    return _truncate(text), None


def _extract_xlsx(raw):
    try:
        import openpyxl
    except ImportError:
        return None, (
            "XLSX support isn't installed on this machine "
            "(pip install openpyxl --break-system-packages)"
        )
    try:
        workbook = openpyxl.load_workbook(io.BytesIO(raw), data_only=True, read_only=True)
    except Exception as e:
        return None, f"couldn't open as an XLSX: {e}"

    sheets_text = []
    MAX_ROWS_PER_SHEET = 500  # generous for a config/data sheet, bounded
    # against someone uploading a 100k-row export and blowing the budget
    # before the overall MAX_EXTRACTED_CHARS truncation even kicks in.
    for sheet in workbook.worksheets:
        rows_out = []
        for row_idx, row in enumerate(sheet.iter_rows(values_only=True)):
            if row_idx >= MAX_ROWS_PER_SHEET:
                rows_out.append(f"[... sheet truncated at {MAX_ROWS_PER_SHEET} rows ...]")
                break
            cells = ["" if v is None else str(v) for v in row]
            if any(c.strip() for c in cells):
                rows_out.append(",".join(cells))
        if rows_out:
            sheets_text.append(f"--- sheet: {sheet.title} ---\n" + "\n".join(rows_out))

    text = "\n\n".join(sheets_text)
    if not text.strip():
        return None, "no data found in this XLSX (it may be empty)"
    return _truncate(text), None


def _extract_csv_like(raw, delimiter=","):
    try:
        decoded = raw.decode("utf-8")
    except UnicodeDecodeError:
        try:
            decoded = raw.decode("latin-1")
        except Exception as e:
            return None, f"couldn't decode as text: {e}"

    try:
        reader = csv.reader(io.StringIO(decoded), delimiter=delimiter)
        rows = [", ".join(cell.strip() for cell in row) for row in reader]
    except Exception as e:
        return None, f"couldn't parse as delimited data: {e}"

    text = "\n".join(r for r in rows if r.strip())
    if not text.strip():
        return None, "file is empty"
    return _truncate(text), None


def _extract_ipynb(raw):
    try:
        notebook = json.loads(raw.decode("utf-8"))
    except (UnicodeDecodeError, json.JSONDecodeError) as e:
        return None, f"couldn't parse as a Jupyter notebook: {e}"

    cells_out = []
    for cell in notebook.get("cells", []):
        source = cell.get("source", "")
        if isinstance(source, list):
            source = "".join(source)
        if not source.strip():
            continue
        cell_type = cell.get("cell_type", "code")
        if cell_type == "markdown":
            cells_out.append(source)
        else:
            cells_out.append(f"```python\n{source}\n```")

    text = "\n\n".join(cells_out)
    if not text.strip():
        return None, "no cells with content found in this notebook"
    return _truncate(text), None


def _extract_odt(raw):
    # ODF documents are a zip of XML — extracted here with stdlib zipfile
    # + a regex strip of tags rather than pulling in odfpy, since the
    # content.xml text-node structure is simple enough that a full ODF
    # library is more dependency weight than this format is likely worth
    # for Perla's use case. Good enough for prose; doesn't preserve
    # tables/formatting.
    try:
        with zipfile.ZipFile(io.BytesIO(raw)) as z:
            xml_bytes = z.read("content.xml")
    except (zipfile.BadZipFile, KeyError) as e:
        return None, f"couldn't open as an ODT: {e}"

    xml_text = xml_bytes.decode("utf-8", errors="replace")
    # Turn paragraph/line-break tags into newlines before stripping all
    # other tags, so the result reads as prose instead of one run-on line.
    xml_text = re.sub(r"</text:p>|<text:line-break/>", "\n", xml_text)
    text = re.sub(r"<[^>]+>", "", xml_text)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text).strip()

    if not text:
        return None, "no extractable text found in this ODT"
    return _truncate(text), None


# ---------------------------------------------------------------------------
# Dispatch tables + public entrypoints
# ---------------------------------------------------------------------------

# Formats extracted as TEXT and inlined into the prompt.
_TEXT_EXTRACTORS = {
    ".docx": _extract_docx,
    ".xlsx": _extract_xlsx,
    ".ipynb": lambda raw: _extract_ipynb(raw),
    ".odt": _extract_odt,
    ".tsv": lambda raw: _extract_csv_like(raw, delimiter="\t"),
}

# Formats rendered as PAGE IMAGES and sent through the vision path instead.
_IMAGE_EXTRACTORS = {
    ".pdf": _rasterize_pdf,
    ".pptx": _rasterize_pptx,
}

# Formats this module CAN target eventually but doesn't yet — kept here so
# perla-companion.py's allowlist and this module's capability list can be
# cross-checked instead of drifting apart silently.
PLANNED_UNSUPPORTED = {".epub", ".doc", ".xls", ".ppt", ".rtf"}

SUPPORTED_EXTENSIONS = frozenset(_TEXT_EXTRACTORS.keys() | _IMAGE_EXTRACTORS.keys())
SUPPORTED_TEXT_EXTENSIONS = frozenset(_TEXT_EXTRACTORS.keys())
SUPPORTED_IMAGE_EXTENSIONS = frozenset(_IMAGE_EXTRACTORS.keys())


def extract_text(raw, filename):
    """Extract plain text from a document format that uses the TEXT
    strategy (see module docstring). Returns (text, error) — None/reason
    on failure. Calling this on a format that uses the IMAGE strategy
    (.pdf/.pptx) returns an error rather than silently doing the wrong
    thing; use rasterize_pages or convert for those instead.
    """
    ext = os.path.splitext(filename or "")[1].lower()
    extractor = _TEXT_EXTRACTORS.get(ext)
    if extractor is None:
        if ext in _IMAGE_EXTRACTORS:
            return None, f"'{filename}' ({ext}) is rendered as images, not extracted as text — use rasterize_pages/convert"
        if ext in PLANNED_UNSUPPORTED:
            return None, f"'{filename}' ({ext}) isn't supported yet"
        return None, f"'{filename}' ({ext or 'no extension'}) is not a supported document type"

    try:
        text, error = extractor(raw)
    except Exception as e:  # noqa: BLE001 — last-resort guard so one
        # malformed upload can never raise out of this module into the
        # daemon's request handler; every extractor above already has
        # its own targeted except clauses, this only catches something
        # genuinely unanticipated.
        return None, f"'{filename}' failed to extract ({e})"

    if error:
        return None, f"'{filename}': {error}"
    return text, None


def rasterize_pages(raw, filename):
    """Render a document format that uses the IMAGE strategy (.pdf/.pptx)
    into a list of PNG byte-blobs, one per page/slide. Returns
    (list_of_png_bytes, error) — None/reason on failure. Calling this on
    a TEXT-strategy format returns an error; use extract_text for those.
    """
    ext = os.path.splitext(filename or "")[1].lower()
    rasterizer = _IMAGE_EXTRACTORS.get(ext)
    if rasterizer is None:
        if ext in _TEXT_EXTRACTORS:
            return None, f"'{filename}' ({ext}) is extracted as text, not rendered as images — use extract_text/convert"
        if ext in PLANNED_UNSUPPORTED:
            return None, f"'{filename}' ({ext}) isn't supported yet"
        return None, f"'{filename}' ({ext or 'no extension'}) is not a supported document type"

    try:
        pages, error = rasterizer(raw)
    except Exception as e:  # noqa: BLE001 — see extract_text's rationale
        return None, f"'{filename}' failed to render ({e})"

    if error:
        return None, f"'{filename}': {error}"
    return pages, None


def convert(raw, filename):
    """Single entrypoint perla-companion.py should call: figures out
    which strategy `filename`'s extension uses and dispatches to it.

    Returns (kind, payload, error):
      - kind == "text":   payload is a str  (or None on error)
      - kind == "images": payload is a list of PNG bytes  (or None on error)
      - kind == None:     unsupported format; payload is None, error is set

    This is the only function perla-companion.py needs to know about —
    it doesn't need its own copy of which extension uses which strategy.
    """
    ext = os.path.splitext(filename or "")[1].lower()
    if ext in _TEXT_EXTRACTORS:
        text, error = extract_text(raw, filename)
        return "text", text, error
    if ext in _IMAGE_EXTRACTORS:
        pages, error = rasterize_pages(raw, filename)
        return "images", pages, error
    if ext in PLANNED_UNSUPPORTED:
        return None, None, f"'{filename}' ({ext}) isn't supported yet"
    return None, None, f"'{filename}' ({ext or 'no extension'}) is not a supported document type"


if __name__ == "__main__":
    import sys
    if len(sys.argv) != 2:
        print("usage: perla-textify.py <file>", file=sys.stderr)
        sys.exit(1)
    path = sys.argv[1]
    with open(path, "rb") as f:
        data = f.read()
    result_kind, result_payload, result_error = convert(data, os.path.basename(path))
    if result_error:
        print(f"ERROR: {result_error}", file=sys.stderr)
        sys.exit(1)
    if result_kind == "text":
        print(result_payload)
    else:
        base = os.path.splitext(path)[0]
        for i, png_bytes in enumerate(result_payload, start=1):
            out_path = f"{base}-page{i}.png"
            with open(out_path, "wb") as f:
                f.write(png_bytes)
            print(f"wrote {out_path}", file=sys.stderr)

